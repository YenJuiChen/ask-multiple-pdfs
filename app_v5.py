import streamlit as st
import requests
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings, HuggingFaceInstructEmbeddings
from langchain.vectorstores import FAISS
from langchain.chat_models import ChatOpenAI
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from htmlTemplates import css, bot_template, user_template
from bs4 import BeautifulSoup
from docx import Document
import docx2txt
import os

# Function to get text from DOCX files

def get_text(uploaded_files):
    text = ""
    for docx in uploaded_files:
        doc = Document(docx)
        for para in doc.paragraphs:
            text += para.text + '\n'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + '\t'
                text += '\n'
        # Extract images
        temp_dir = '/tmp/docx_images'
        os.makedirs(temp_dir, exist_ok=True)
        docx2txt.process(docx, temp_dir)
        for img_file in os.listdir(temp_dir):
            with open(os.path.join(temp_dir, img_file), 'rb') as f:
                img_data = f.read()
            # TODO: Add image data to FAISS vector store or other processing

    for pdf in uploaded_files:
        try:
            pdf_reader = PdfReader(pdf)
            for page in pdf_reader.pages:
                text += page.extract_text()
        except Exception as e:
            st.write(f"Error reading PDF: {e}")
    return text


def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=500,
        chunk_overlap=100,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks


def get_vectorstore(text_chunks):
    if not text_chunks:
        st.write("No text chunks available for vectorization.")
        return None
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_texts(texts=text_chunks, embedding=embeddings)
    return vectorstore

def get_conversation_chain(vectorstore):
    llm = ChatOpenAI(model_name='gpt-4')
    #llm = HuggingFaceHub(repo_id="google/flan-t5-xxl", model_kwargs={"temperature":0.8,"max_length":512})

    memory = ConversationBufferMemory(
        memory_key='chat_history', return_messages=True)
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        memory=memory
    )
    return conversation_chain


def handle_userinput(user_question):
    if st.session_state.conversation is None:
        st.write("對話系統還未初始化。請先上傳PDF文件並點擊 '運行'。")
        return

    response = st.session_state.conversation({'question': user_question})
    st.session_state.chat_history = response['chat_history']

    # 檢查response是否包含有效答案
    if not response['chat_history']:
        # 進行網路搜尋
        search_url = f"https://www.google.com/search?q={user_question}"
        search_response = requests.get(search_url)
        
        if search_response.status_code == 200:
            soup = BeautifulSoup(search_response.text, 'html.parser')
            search_results = ' '.join([result.text for result in soup.find_all('p')])
            
            # 將搜尋結果加入向量資料庫
            text_chunks = get_text_chunks(search_results)
            st.session_state.conversation.vectorstore.add(text_chunks)
            
            # 重新回答問題
            response = st.session_state.conversation({'question': user_question})
            st.session_state.chat_history = response['chat_history']
            
            # 如果仍然没有答案，显示搜索结果
            if not response['chat_history']:
                st.write(f"對於您的問題，我在網上找到了以下資訊：{search_results}")
                return

    for i, message in enumerate(st.session_state.chat_history):
        if i % 2 == 0:
            st.write(user_template.replace(
                "{{MSG}}", message.content), unsafe_allow_html=True)
        else:
            st.write(bot_template.replace(
                "{{MSG}}", message.content), unsafe_allow_html=True)


def main():
    load_dotenv()
    st.set_page_config(page_title="YenJuiChen_PDF專家系統_v4",
                       page_icon="\u2600")
    st.write(css, unsafe_allow_html=True)

    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []  # Initialize chat_history

    st.header("YenJuiChen_PDF專家系統_v4 \u2600")
    user_question = st.text_input("(Step 3)輸入問題:")
    if user_question:
        handle_userinput(user_question)

    with st.sidebar:
        st.subheader("你的檔案")
        uploaded_files = st.file_uploader(
            "(Step 1)上傳PDF檔案並點擊 '運行'", accept_multiple_files=True, type=['pdf', 'docx'])
        if st.button("(Step 2)運行"):
            with st.spinner("運行中"):
                # get text
                raw_text = get_text(uploaded_files)

                # get the text chunks
                text_chunks = get_text_chunks(raw_text)

                # create vector store
                vectorstore = get_vectorstore(text_chunks)

                if vectorstore is not None:
                   # create conversation chain
                   st.session_state.conversation = get_conversation_chain(vectorstore)


if __name__ == '__main__':
    main()
